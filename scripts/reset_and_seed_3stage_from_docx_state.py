# scripts/reset_and_seed_3stage_from_docx_state.py
"""
Reset exam data and seed a 3-stage quiz using questions parsed from a .docx file.

Requirements:
    pip install python-docx

Run:
    python manage.py shell -c "from scripts.reset_and_seed_3stage_from_docx_state import run; run()"

Notes:
- Expects your Question/QuestionOption schema like in the previous script.
- Parses blocks in the docx that look like:
      Question: <text>
      (a) option text
      (b) option text
      ...
      Answer: a
- Difficulty distribution is applied 40% EASY, 40% MEDIUM, 20% HARD across the extracted questions.
- Stage 2 & 3 are gated (requires_admission=True). See the "STATE-BASED ADMISSIONS" section for examples.

If your doc path differs, change DOCX_PATH below.
"""

from datetime import timedelta
import re
from django.utils import timezone
from django.db import transaction
from django.db.models import Q

# Third-party
from docx import Document

# Project enums & models (same imports you used before)
from common.enums import Difficulty  # EASY, MEDIUM, HARD
from exams.models import (
    Question, QuestionOption,
    Quiz, QuizStage, StageRandomRule,
    StageQuestion,
    QuizAttempt, QuizStageAttempt, AttemptAnswer, StageAttemptItem,
    LeaderboardEntry, StageAdmission, AccessToken,
)

# If you need to use User/IndianState for admissions
from accounts.models import User  # your custom user
# from accounts.models import IndianState  # if IndianState is declared alongside User, import it accordingly

# -------------------
# CONFIG
# -------------------

# Change this to where your file actually lives in your project/server:
DOCX_PATH = "plab test ..docx"  # e.g. "<project_root>/data/plab test ..docx"
QUIZ_SLUG = "quiz-3stage-from-docx"
QUIZ_TITLE = "Tech Stack Quiz (from Word bank; 3 stages, 20m each)"
QUIZ_DESC = "Imported MCQs from .docx; 3 stages with random selection."

# Stage question counts
S1_COUNT, S2_COUNT, S3_COUNT = 12, 13, 15

# Difficulty distribution over the entire extracted set: (easy, medium, hard)
DIFF_SPLITS = (0.40, 0.40, 0.20)  # 40/40/20 like previous 24/24/12 over 60


# -------------------
# DOCX PARSING
# -------------------

# -------------------
# DOCX PARSING
# -------------------

from docx import Document
import re

_Q_START_RE = re.compile(r"question", re.IGNORECASE)
_OPT_RE     = re.compile(r"^\s*\(?([a-eA-E])\)?[.)]?\s*(.+)$")
_ANS_RE     = re.compile(r"answer", re.IGNORECASE)

def _parse_docx(path):
    doc = Document(path)
    lines = []

    # Collect text from paragraphs
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            lines.append(p.text.strip())

    # Collect text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text and p.text.strip():
                        lines.append(p.text.strip())

    # --- parse lines into blocks ---
    blocks, cur_q = [], None

    def commit():
        nonlocal cur_q
        if cur_q and cur_q.get("options") and cur_q.get("answer_idx") is not None:
            blocks.append(cur_q)
        cur_q = None

    for line in lines:
        low = line.lower()
        if low.startswith("question"):
            commit()
            qtext = line.split(":", 1)[-1].strip() if ":" in line else line
            cur_q = {"text": qtext, "options": [], "answer_idx": None}
        elif re.match(r"^\s*\(?[a-eA-E][).]?\s+", line) and cur_q:
            cur_q["options"].append(line.split(" ", 1)[-1].strip())
        elif low.startswith("answer") and cur_q:
            ans_letter = line.split(":")[-1].strip().lower()[:1]
            cur_q["answer_idx"] = ord(ans_letter) - ord("a")
    commit()

    return blocks


# -------------------
# HELPERS
# -------------------

def _mk_q_from_block(topic, diff, block):
    """
    Create Question + options from a parsed block.
    :param block: {"text": str, "options": [str,...], "answer_idx": int}
    """
    q = Question.objects.create(
        text=block["text"].strip(),
        explanation="",
        question_type="single",
        subspecialty=topic.lower(),  # simple tag; change as you like
        difficulty=diff,
        region_hint="",
        marks=1,
        negative_marks=0.25,
        time_limit_seconds=60,
        is_active=True,
        tags={"source": "docx", "topic": topic.lower()},
    )

    bulk = []
    for i, opt in enumerate(block["options"], start=1):
        is_correct = (i - 1) == int(block["answer_idx"])
        bulk.append(QuestionOption(
            question=q,
            text=opt,
            is_correct=is_correct,
            order=i
        ))
    QuestionOption.objects.bulk_create(bulk)

    return q


def _assign_difficulties(n, splits=DIFF_SPLITS):
    """
    Return a list of Difficulty enums of length n according to splits.
    """
    easy_n  = int(round(n * splits[0]))
    med_n   = int(round(n * splits[1]))
    # ensure total sums to n
    hard_n  = max(0, n - easy_n - med_n)

    diffs = ([Difficulty.EASY] * easy_n) + \
            ([Difficulty.MEDIUM] * med_n) + \
            ([Difficulty.HARD] * hard_n)
    # If rounding caused mismatch, truncate/pad carefully
    if len(diffs) > n:
        diffs = diffs[:n]
    elif len(diffs) < n:
        diffs += [Difficulty.MEDIUM] * (n - len(diffs))
    return diffs


# -------------------
# MAIN seeding routine
# -------------------

def run():
    print("=== RESET (exams-only) → seed 3-stage quiz (20m stages & 20m breaks) FROM DOCX ===")
    now = timezone.now()

    # 0) Wipe exam runtime/config tables AND question bank
    with transaction.atomic():
        AttemptAnswer.objects.all().delete()
        StageAttemptItem.objects.all().delete()
        QuizStageAttempt.objects.all().delete()
        QuizAttempt.objects.all().delete()

        LeaderboardEntry.objects.all().delete()
        StageAdmission.objects.all().delete()
        StageQuestion.objects.all().delete()
        StageRandomRule.objects.all().delete()
        QuizStage.objects.all().delete()
        AccessToken.objects.all().delete()
        Quiz.objects.all().delete()

        QuestionOption.objects.all().delete()
        Question.objects.all().delete()

    print("Cleared exams data + question bank (kept users, anti-cheat logs, etc.).")

    # 1) Parse docx & seed questions
    blocks = _parse_docx(DOCX_PATH)
    if not blocks:
        raise RuntimeError(f"No questions parsed from DOCX at: {DOCX_PATH}")

    diffs = _assign_difficulties(len(blocks))
    # You can set a single topic label; or infer from sections
    topic = "general"

    for i, block in enumerate(blocks):
        _mk_q_from_block(topic, diffs[i], block)

    bank_size = Question.objects.filter(is_active=True).count()
    print("Bank size (active):", bank_size)

    # Compute totals for quiz meta (sum by difficulty)
    easy_total  = Question.objects.filter(difficulty=Difficulty.EASY, is_active=True).count()
    med_total   = Question.objects.filter(difficulty=Difficulty.MEDIUM, is_active=True).count()
    hard_total  = Question.objects.filter(difficulty=Difficulty.HARD, is_active=True).count()

    s1_start = now
    s1_end   = now + timedelta(minutes=20)

    s2_start = now + timedelta(minutes=40)   # 20m break after S1
    s2_end   = now + timedelta(minutes=60)

    s3_start = now + timedelta(minutes=80)   # 20m break after S2
    s3_end   = now + timedelta(minutes=100)

    quiz, _ = Quiz.objects.update_or_create(
        slug=QUIZ_SLUG,
        defaults=dict(
            title=QUIZ_TITLE,
            description=QUIZ_DESC,
            subspecialty="general",
            easy_count=easy_total, medium_count=med_total, hard_count=hard_total,
            start_at=s1_start,
            end_at=s3_end + timedelta(minutes=10),
            duration_seconds=60*120,  # overall cap; per-stage enforced on stages
            pass_threshold_percent=60,
            max_attempts_per_user=1,
            question_count=easy_total + med_total + hard_total,
            shuffle_questions=True,
            shuffle_options=True,
            require_fullscreen=True,
            lock_on_tab_switch=True,
            prerequisite_tutorial=None,
            is_active=True,
        )
    )

    # ----- Stages (same sizing as your earlier script)
    s1, _ = QuizStage.objects.update_or_create(
        quiz=quiz, order=1,
        defaults=dict(
            title="Stage 1", description="Round 1",
            start_at=s1_start, end_at=s1_end,
            duration_seconds=20*60,
            question_count=S1_COUNT,
            shuffle_questions=None, shuffle_options=None,
            is_current=True,
            requires_admission=False,
        )
    )
    s2, _ = QuizStage.objects.update_or_create(
        quiz=quiz, order=2,
        defaults=dict(
            title="Stage 2", description="Round 2",
            start_at=s2_start, end_at=s2_end,
            duration_seconds=20*60,
            question_count=S2_COUNT,
            shuffle_questions=None, shuffle_options=None,
            is_current=False,
            requires_admission=True,   # gated
        )
    )
    s3, _ = QuizStage.objects.update_or_create(
        quiz=quiz, order=3,
        defaults=dict(
            title="Stage 3", description="Round 3",
            start_at=s3_start, end_at=s3_end,
            duration_seconds=20*60,
            question_count=S3_COUNT,
            shuffle_questions=None, shuffle_options=None,
            is_current=False,
            requires_admission=True,   # gated
        )
    )

    # Ensure only Stage 1 is current
    QuizStage.objects.filter(quiz=quiz).exclude(pk=s1.pk).update(is_current=False)

    # Random rules: full bank; counts = stage.question_count
    for st in (s1, s2, s3):
        StageRandomRule.objects.update_or_create(
            stage=st,
            defaults=dict(count=st.question_count, tags_any=[], difficulties=[], subspecialties=[], region_hints=[])
        )

    # Make this the ONLY active quiz
    Quiz.objects.exclude(pk=quiz.pk).update(is_active=False)

    # ----------------------------
    # STATE-BASED ADMISSIONS (optional)
    # ----------------------------
    # By default you might want to admit all verified users to Stage 2 & 3:
    verified_users = User.objects.filter(is_verified=True).only("id")
    StageAdmission.objects.bulk_create([
        StageAdmission(stage=s2, user=u) for u in verified_users
    ] + [
        StageAdmission(stage=s3, user=u) for u in verified_users
    ], ignore_conflicts=True)

    # If you instead want to admit ONLY users of certain Indian states:
    # from accounts.models import IndianState
    # allowed_states = [IndianState.MAHARASHTRA, IndianState.KARNATAKA]
    # state_users = User.objects.filter(is_verified=True, state__in=allowed_states).only("id")
    # StageAdmission.objects.bulk_create(
    #     [StageAdmission(stage=s2, user=u) for u in state_users] +
    #     [StageAdmission(stage=s3, user=u) for u in state_users],
    #     ignore_conflicts=True
    # )

    print("Active quiz:", quiz.slug)
    print("Stages:",
          list(QuizStage.objects.filter(quiz=quiz).order_by("order")
               .values("order","title","start_at","end_at","question_count",
                       "duration_seconds","is_current","requires_admission")))
    print("=== DONE ===")
